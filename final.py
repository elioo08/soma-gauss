import PySimpleGUI as sg
import docx 
from docx.shared import Inches
from docx.shared import Pt 

layout = [[sg.Text('Soma de Sauss')],      
                 [sg.Text('Primeiro termo'), sg.InputText(key='ia1')],   
                 [sg.Text('Ultimo termo'), sg.InputText(key='ian')],  
                 [sg.Text('Espaço entre 2 concecutivos'), sg.InputText(key='ie')],
                 [sg.Submit('Calcular'), sg.Cancel()]]      

window = sg.Window('Soma de Gauss', layout)    

event, values = window.read()  

if event == 'Calcular':
    a1 = float(values['ia1'])
    an = float(values['ian'])
    e = float(values['ie'])
    n = an - a1 + e

    s2 = n*(a1 + an)
    s1 = s2 / 2

    doc = docx.Document() 
    doc.add_heading('Soma de Gauss', 0) # Imprime o cabeçalho
    doc.add_heading('Resultado', 1)     # Imprime o resultado 
    para = doc.add_paragraph().add_run(f'{s1:.0f}') 
    para.font.size = Pt(15)  
    doc.add_heading('Contas', 1) 
    doc.add_heading('Fórmula:', 2) 
    doc.add_picture('img1-doc.png', width = Inches(1.50))
    doc.add_heading('Aplicando na sua conta:', 3) 
    doc.add_paragraph(f'a1 = {a1} \n an = {an} \n n = {n}') 
    doc.add_paragraph(f'S2 = {n} . ({a1 + an}) = {s2}') 
    doc.add_paragraph(f's1 = S2 / 2 = {s1}') 
    resultado = doc.add_paragraph(f'Chegando no resultado: {s1}') 
    resultado.bold = True
    doc.save('resultado.docx')

    sg.popup_ok(f'Resultado: {s1:.0f} \n Veja mais no documento gerado!')

    print(f'Resultado: {s1:.0f}')
    print("Veja mais no documento!! ^__^")

window.close()

